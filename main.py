import os
import streamlit as st
import io
from docx import Document
from datetime import datetime
from openai import OpenAI
from dotenv import load_dotenv

# --- 页面全局配置 (必须放在首行) ---
st.set_page_config(page_title='Kznk的智能总结引擎', page_icon='🚀', layout='wide')

# --- 导入 PDF 处理库 ---
try:
    import fitz  # PyMuPDF
except ImportError:
    fitz = None

# 加载 .env 文件中的配置
load_dotenv()

# 初始化 OpenAI 客户端
client = OpenAI(
    api_key=os.getenv("OPENAI_API_KEY"),
    base_url=os.getenv("OPENAI_BASE_URL", "https://api.openai.com/v1")
)

def read_pdf_file(uploaded_file):
    """提取 PDF 文件中的文字 (支持 Streamlit 上传对象)"""
    if fitz is None:
        st.error("❌ 错误：尚未安装 PyMuPDF 库。请在终端运行: pip install pymupdf")
        return None
    
    try:
        # 将上传的文件流读入内存
        file_bytes = uploaded_file.read()
        text = ""
        # 使用 stream 参数打开 PDF
        with fitz.open(stream=file_bytes, filetype="pdf") as doc:
            for page in doc:
                text += page.get_text()
        return text
    except Exception as e:
        st.error(f"❌ 读取 PDF 文件时发生错误: {e}")
        return None

def summarize_text(text, target_lang, target_len, target_style):
    """调用 API 进行个性化配置的深度总结"""
    if not text or len(text.strip()) == 0:
        st.warning("⚠️ 警告：提取到的文本内容为空，无法生成总结。")
        return None

    try:
        response = client.chat.completions.create(
            model="deepseek-chat",
            messages=[
                {
                    "role": "system", 
                    "content": (
                        f"你是一个极具专业水平的文件总结专家。请对用户提供的文本进行深度总结。\n"
                        f"请务必遵守以下个性化指令：\n"
                        f"1. 输出语言：{target_lang}\n"
                        f"2. 总结长度：目标约为 {target_len}\n"
                        f"3. 语气风格：采用 {target_style} 的笔触\n"
                        f"4. 结构化输出：必须包含【核心观点提取】和【详细内容展开】两个部分。\n"
                        f"5. 格式规范：使用清晰的分段、列表（Bullet Points）或数字编号。"
                    )
                },
                {"role": "user", "content": f"请针对以下文本内容进行深度总结：\n\n{text}"}
            ]
        )
        return response.choices[0].message.content
    except Exception as e:
        st.error(f"❌ 调用 API 时发生错误: {e}")
        return None

def main():
    # --- 初始化记忆中枢 (Session State) ---
    if "messages" not in st.session_state:
        st.session_state.messages = []
    if "doc_text" not in st.session_state:
        st.session_state.doc_text = ""

    # --- 主界面标题与美化 ---
    st.markdown("<h1 style='text-align: center; color: #1E90FF;'>✨ Kznk 智能文档分析核心</h1>", unsafe_allow_html=True)
    st.divider()
    st.markdown("<p style='text-align: center; font-style: italic; color: #666;'>欢迎，指挥官。请在下方开启您的智慧分析之旅。</p>", unsafe_allow_html=True)
    
    # --- 侧边栏：安全校验 ---
    with st.sidebar:
        st.markdown("### 🔐 安全访问")
        # 增加口令输入框
        password = st.text_input("请输入访问口令", type="password")
        
        st.divider()

    # --- 逻辑拦截 ---
    # 设定访问口令
    SECRET_PASSWORD = "2026"

    if password != SECRET_PASSWORD:
        # 如果口令不正确，显示警示信息并停止后续渲染
        st.warning("🛡️ 请在左侧侧边栏输入正确口令以解锁功能")
        st.info("提示：请与系统管理员联系获取访问口令。")
        
        with st.sidebar:
            st.header("关于助手")
            st.write("本助手基于 DeepSeek API 开发，支持多种格式文档的一键总结与对话。")
            st.divider()
            st.caption("版本: 8.0 (ChatPDF V2.0)")
        return # 提前结束，不显示下方组件

    # --- 侧边栏：交互控制台 (仅在口令正确时显示) ---
    with st.sidebar:
        st.header("⚙️ 总结设置")
        
        # 1. 输出语言
        out_lang = st.selectbox("输出语言", ["中文", "English", "日本語"])
        
        # 2. 总结长度
        out_len = st.selectbox("总结长度", [
            "简短摘要 (约100字)", 
            "标准总结 (约300字)", 
            "深度解析 (约800字)"
        ])
        
        # 3. 语气风格
        out_style = st.selectbox("语气风格", ["专业严谨", "幽默风趣", "大白话讲故事"])
        
        st.divider()
        st.button("🧹 清除聊天记录", on_click=lambda: st.session_state.messages.clear())
        
        st.header("关于助手")
        st.write("本助手基于 DeepSeek API 开发，支持多种格式文档的一键总结与对话。")
        st.divider()
        st.caption("版本: 8.0 (ChatPDF V2.0)")

    # --- 核心功能区 (口令正确时) ---
    st.info(f"✅ 访问已解锁！配置：{out_lang} | {out_len} | {out_style}")

    # 检查 API KEY
    if not os.getenv("OPENAI_API_KEY"):
        st.error("💡 请先在根目录的 `.env` 文件中设置您的 `OPENAI_API_KEY` 以继续。")
        st.stop()

    # 文件上传组件
    uploaded_file = st.file_uploader("选择一个文件", type=["txt", "md", "pdf"])

    if uploaded_file is not None:
        file_details = {"文件名": uploaded_file.name, "文件大小": f"{uploaded_file.size / 1024:.2f} KB"}
        st.write("📁 文件已选择:", file_details)

        # ---------------- 总结功能 ----------------
        if st.button("✨ 开始总结", type="primary"):
            with st.spinner("🔍 正在根据您的个性化偏好进行总结，请稍候..."):
                content = ""
                # 根据类型读取内容
                if uploaded_file.name.endswith(".pdf"):
                    content = read_pdf_file(uploaded_file)
                else:
                    # 读取文本文件
                    content = uploaded_file.read().decode("utf-8")

                if content:
                    # 保存文档上下文
                    st.session_state.doc_text = content
                    
                    summary = summarize_text(content, out_lang, out_len, out_style)
                    
                    if summary:
                        st.success("✅ 总结完成！")
                        st.markdown("### 📝 总结结果：")
                        # 使用二级容器展示结果，看起来更精美
                        st.markdown(f"> **生成时间**：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
                        st.markdown(summary)

                        # 准备下载内容
                        download_content = (
                            f"【智能总结报告】\n"
                            f"{'='*30}\n"
                            f"文件名：{uploaded_file.name}\n"
                            f"配置：语言={out_lang}, 长度={out_len}, 风格={out_style}\n"
                            f"生成时间：{datetime.now()}\n"
                            f"{'='*30}\n\n"
                            f"{summary}"
                        )
                        
                        st.download_button(
                            label="📥 下载总结结果 (TXT)",
                            data=download_content,
                            file_name=f"{os.path.splitext(uploaded_file.name)[0]}_深度解析.txt",
                            mime="text/plain"
                        )

                        # --- 新增：Word 导出功能 (V3.0) ---
                        doc = Document()
                        doc.add_heading(f"智能总结报告", 0)
                        doc.add_paragraph(f"文件名：{uploaded_file.name}")
                        doc.add_paragraph(f"配置：语言={out_lang}, 长度={out_len}, 风格={out_style}")
                        doc.add_paragraph(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
                        doc.add_section()
                        doc.add_paragraph(summary)

                        # 保存到内存字节流
                        doc_io = io.BytesIO()
                        doc.save(doc_io)
                        doc_io.seek(0)

                        st.download_button(
                            label="📥 一键下载总结报告 (Word)",
                            data=doc_io,
                            file_name=f"{os.path.splitext(uploaded_file.name)[0]}_深度解析.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )

        # ---------------- 对话功能 (ChatPDF 模式) ----------------
        st.divider()
        st.subheader("💬 文档对话助手")
        
        if not st.session_state.doc_text:
            st.info("💡 请先点击上面的“开始总结”按钮解析文档，解析后即可开启对话问答。")
        else:
            # 渲染历史记录
            for message in st.session_state.messages:
                with st.chat_message(message["role"]):
                    st.markdown(message["content"])

            # 接受用户提问
            if prompt := st.chat_input("向 AI 提问关于这份文档的内容..."):
                # 记录并显示用户问题
                st.session_state.messages.append({"role": "user", "content": prompt})
                with st.chat_message("user"):
                    st.markdown(prompt)

                # 生成 AI 回答
                with st.chat_message("assistant"):
                    with st.spinner("🧠 正在检索文档内容并思考..."):
                        try:
                            response = client.chat.completions.create(
                                model="deepseek-chat",
                                messages=[
                                    {
                                        "role": "system", 
                                        "content": (
                                            "你是一个文档问答助手，请严格根据以下提供的文档内容回答用户的问题。\n"
                                            "1. 如果文档中没有相关信息，请直接回答'抱歉，在文档中未找到相关内容'。\n"
                                            "2. 你的回答应准确、简洁、专业。\n"
                                            f"文档内容：\n{st.session_state.doc_text[:10000]}" # 限制长度防止超限
                                        )
                                    },
                                    *st.session_state.messages[-10:] # 只携带最近10轮对话，保持记忆并节省 tokens
                                ]
                            )
                            answer = response.choices[0].message.content
                            st.markdown(answer)
                            # 保存助手回答
                            st.session_state.messages.append({"role": "assistant", "content": answer})
                        except Exception as e:
                            st.error(f"❌ 对话发生错误: {e}")

if __name__ == "__main__":
    main()
